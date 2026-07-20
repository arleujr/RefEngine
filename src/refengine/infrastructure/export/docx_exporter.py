from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from refengine.domain.models import ProcessedDocument
from refengine.services.reference_formatter import ReferenceFormatter


def export_references_docx(
    documents: list[ProcessedDocument],
    output_path: Path,
    *,
    heading_text: str = "REFERÊNCIAS",
    notice: str | None = None,
) -> None:
    """Create a copy-ready reference list without touching the user's manuscript."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.bottom_margin = Cm(2)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(24)
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run(heading_text)
    _format_run(run, bold=True)

    if notice:
        notice_paragraph = document.add_paragraph()
        notice_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        notice_paragraph.paragraph_format.space_after = Pt(18)
        notice_run = notice_paragraph.add_run(notice)
        _format_run(notice_run, bold=True)

    for item in documents:
        if not item.generated_reference:
            continue
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.space_after = Pt(12)
        reference = item.generated_reference
        resolved = item.resolved_bibliography
        emphasis_values = (
            ReferenceFormatter().emphasis_values(resolved) if resolved is not None else []
        )
        _append_reference_runs(paragraph, reference, emphasis_values)

    document.core_properties.title = "Referências UFV/ABNT — RefEngine"
    document.core_properties.subject = "Lista de referências gerada localmente"
    document.core_properties.author = "Arleu Júnior"
    document.save(str(output_path))


def _append_reference_runs(
    paragraph: Paragraph, reference: str, emphasis_values: list[str]
) -> None:
    """Write reference text while emphasizing catalog-selected title segments."""
    remaining = reference
    values = [value for value in emphasis_values if value and value in reference]
    if not values:
        _format_run(paragraph.add_run(reference))
        return

    while remaining:
        matches = [(remaining.find(value), value) for value in values if remaining.find(value) >= 0]
        if not matches:
            _format_run(paragraph.add_run(remaining))
            break
        index, value = min(matches, key=lambda item: item[0])
        if index:
            _format_run(paragraph.add_run(remaining[:index]))
        _format_run(paragraph.add_run(value), bold=True)
        remaining = remaining[index + len(value) :]


def _format_run(run: Run, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
